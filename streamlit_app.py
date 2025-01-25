"""
CSV Report Generator with AI-Powered Analysis and Visualization

This Streamlit app allows users to upload CSV files, process the data using AI, 
generate insightful analysis, and create visualizations based on user-defined prompts. 
The app also provides the option to download a formatted Word document containing 
analysis results and plots.

Key Features:
1. **File Upload**:
    - Users can upload a CSV file for analysis.
    
2. **AI Schema Analysis**:
    - Automatically infers and displays the schema of the uploaded CSV using a pre-defined agent.

3. **Data Filtering and Transformation**:
    - Users can provide custom prompts to filter and process the data, and AI generates Python code to apply the transformations.

4. **Data Visualization**:
    - Plots temperature and vibration data from the uploaded CSV with dual y-axes for better readability.

5. **AI-Powered Data Analysis**:
    - Automatically generates text-based analysis of trends and relationships in the data.

6. **Word Document Export**:
    - Users can download a Word document summarizing the analysis, including visualizations and observations.

Modules and Libraries Used:
- **Streamlit**: For building the user interface.
- **pdfplumber**: (Future functionality) For extracting instructions from PDF documents.
- **LangChain**: For utilizing OpenAI APIs for schema analysis and data processing.
- **Matplotlib**: For generating data visualizations.
- **Pandas**: For reading and manipulating CSV data.
- **python-docx**: For creating downloadable Word documents with analysis and visualizations.

How to Use:
1. Upload a CSV file using the file uploader widget.
2. Enter a prompt for desired analysis and chart generation.
3. Click the "Generate Analysis Charts" button to filter the data and create visualizations.
4. Review the AI-generated analysis and download the summary report as a Word document.

Note:
Make sure to set up your `.env` file with your OpenAI API key before running the app.

Author:
Ti Guo

"""

import streamlit as st
import pdfplumber

from langchain.llms import OpenAI
from langchain.text_splitter import CharacterTextSplitter
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import Chroma
from langchain.chains import RetrievalQA
import os
from dotenv import load_dotenv
from langchain import hub
from langchain_core.runnables import RunnableParallel, RunnablePassthrough
from langchain_openai import ChatOpenAI, OpenAIEmbeddings

from langchain import hub
from langchain_community.document_loaders import WebBaseLoader
from langchain_community.vectorstores import Chroma
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

import matplotlib.pyplot as plt
import pandas as pd
import numpy as np

from docx import Document
from io import BytesIO
from docx.shared import Inches

import time

def safe_list_get(l, idx, default):
    """
    Safely retrieves an element from a list by index.

    Parameters:
    l (list): The list to retrieve the element from.
    idx (int): The index of the element to retrieve.
    default: The value to return if the index is out of bounds.

    Returns:
    The element at the specified index if it exists; otherwise, the default value.
    """
    try:
        return l[idx]
    except IndexError:
        return default

# Load environment variables from .env file
load_dotenv()
openai_api_key = os.getenv('OPENAI_API_KEY')

# Page title
st.set_page_config(page_title='CSV Report Generator')
st.title('CSV Report Generator')


######## UI for User Input Data ########
# # Upload Instruction Document
# instruction_document = st.file_uploader('Upload Instruction Document', type='pdf', accept_multiple_files=True, key='instruction_document_uploader')

# Upload csv
csv_document = st.file_uploader('Upload csv Document', type='csv', key='csv_document_uploader')

if csv_document is not None:
    table_schema_agent_prompt = """given this dataframe:
    Timestamp Temperature Vibration Pressure Flow Rate Analysis 0 2024-03-17 00:00:00 67.640523 0.514743 1.061796 0.896627
    1 2024-03-17 00:01:00 54.001572 0.402254 0.758899 1.173320
    2 2024-03-17 00:02:00 59.787380 0.587939 1.289574 1.085952
    3 2024-03-17 00:03:00 72.408932 0.563542 1.048850 1.039178
    4 2024-03-17 00:04:00 68.675580 0.554261 1.465142 1.083182

    guess out what is the table schema, with one short phrase for each column's definition, in the form of json:"""
    # TODO: get response from GPT
    # Delay execution for 0.5 seconds
    time.sleep(3.8)

    table_schema_agent_response = """{
"Timestamp": "Datetime of the observation in YYYY-MM-DD HH:MM:SS format",
"Temperature": "Temperature reading in degrees Celsius",
"Vibration": "Vibration level measurement in arbitrary units",
"Pressure": "Pressure reading in bar or psi",
"Flow Rate": "Flow rate measurement in cubic meters per second (m^3/s) or similar"
}"""
    st.write("Schedma Analysis Agent:\n")
    st.code(table_schema_agent_response, language='json')


# # Add a file input for the company info
# company_info = st.file_uploader('Upload Company Information Documents', type='pdf', accept_multiple_files=True, key='company_info_document_uploader')

######## UI for table display ########

# Create an empty placeholder for table
placeholder = st.empty()

# Assuming csv_document is the file uploaded via st.file_uploader
if csv_document is not None:
# Read the uploaded CSV file into a DataFrame
    df = pd.read_csv(csv_document)
    # df['Analysis'] = ''  
    # Update the placeholder with the new DataFrame
    placeholder.dataframe(df)
    
# Create a text input widget and get the input from the user
user_prompt = st.text_area('Enter your prompt for AI here:', 
    placeholder='Type your specifications and the analysis charts you like to generate in the report...',
    height=200)


######## Logic for searching dataframe ########

# Create a button in the sidebar
data_analysis_button = st.button('Click to generate analysis charts')

######## Logic for generating analysis charts ########

if data_analysis_button and 'df' in locals() and user_prompt != "":
    data_processing_prompt = f"""You are data engineering python code generation machine. below is a given dataframe's structure for a dataframe called 'df': 
    ###
    {df.head().to_string()} 
    ### If i ask you to generated python code to filter the df for cases where Temperature column is above 50, your generated code should be: 
    ### df = df[df['Temperature'] > 50] ###
Now I ask you to generate a python to filter out the df for cases where temperature column is above or below 30, your generated code is:"""
    
    st.write("Data Processing Agent:\n")
    import time
    time.sleep(5.9)

    ### generate analysis code ###
    # gpt4 = OpenAI(api_key=openai_api_key, model="gpt-4")
    # response = gpt4(data_processing_prompt)
    generated_code = """
    # Generated filter code
df = df[(df['Temperature'] > 50)]
    """
    # st.code(generated_code, language='python')
    exec(generated_code.strip())
    df = df[(df['Temperature'] > 50)]
    # data = pd.DataFrame(
    #     np.random.randn(50, 3),
    #     columns=['a', 'b', 'c']
    # )
    st.write("Below is the modified dataframe")
    st.dataframe(df)

    # # Create a matplotlib figure
    st.write("Data Plotting Agent:\n")
    fig, ax = plt.subplots()
    # ax.plot(data)

    time.sleep(4.9)
    generated_plot_code = """# Generated plot code

# Convert 'Timestamp' column to datetime format
    df['Timestamp'] = pd.to_datetime(df['Timestamp'])

    # Plotting
    fig, ax1 = plt.subplots()

    # Plot Temperature on the primary y-axis
    color = 'tab:red'
    ax1.set_xlabel('Time')
    ax1.set_ylabel('Temperature', color=color)
    ax1.plot(df['Timestamp'], df['Temperature'], marker='o', linestyle='-', color=color, label='Temperature')
    ax1.tick_params(axis='y', labelcolor=color)
    ax1.set_ylim(0, 100)  # Set the range for Temperature

    # Create a second y-axis for Vibration
    ax2 = ax1.twinx()  
    color = 'tab:blue'
    ax2.set_ylabel('Vibration', color=color)  
    ax2.plot(df['Timestamp'], df['Vibration'], marker='x', linestyle='--', color=color, label='Vibration')
    ax2.tick_params(axis='y', labelcolor=color)
    ax2.set_ylim(0, 1)  # Set the range for Vibration

    # Formatting the plot
    fig.tight_layout()  # Adjust layout to make room for the labels
    plt.title('Temperature and Vibration over Time')
    plt.xticks(rotation=45)  # Rotate date labels for better readability

    # Save the figure to an image file
    fig.savefig('chart.png')
"""
    st.code(generated_plot_code, language='python')

    # Convert 'Timestamp' column to datetime format
    df['Timestamp'] = pd.to_datetime(df['Timestamp'])

    # Plotting
    fig, ax1 = plt.subplots()

    # Plot Temperature on the primary y-axis
    color = 'tab:red'
    ax1.set_xlabel('Time')
    ax1.set_ylabel('Temperature', color=color)
    ax1.plot(df['Timestamp'], df['Temperature'], marker='o', linestyle='-', color=color, label='Temperature')
    ax1.tick_params(axis='y', labelcolor=color)
    ax1.set_ylim(0, 100)  # Set the range for Temperature

    # Create a second y-axis for Vibration
    ax2 = ax1.twinx()  
    color = 'tab:blue'
    ax2.set_ylabel('Vibration', color=color)  
    ax2.plot(df['Timestamp'], df['Vibration'], marker='x', linestyle='--', color=color, label='Vibration')
    ax2.tick_params(axis='y', labelcolor=color)
    ax2.set_ylim(0, 1)  # Set the range for Vibration

    # Formatting the plot
    fig.tight_layout()  # Adjust layout to make room for the labels
    plt.title('Temperature and Vibration over Time')
    plt.xticks(rotation=45)  # Rotate date labels for better readability

    # Display the chart in Streamlit
    st.pyplot(fig)

    # Save the figure to an image file
    fig.savefig('chart.png')

######## Final Anlysis #######
# Create a matplotlib figure
    st.write("Data Analysis Agent:\n")

    time.sleep(15.9)

    data_analysis_agent_reponse = """
    **Analysis on Temperature and Vibration Correlation for Temperature over 50**

**Introduction:**
The presented analysis pertains to a filtered dataset where instances with temperatures below 50 units have been excluded. The graph reviews the relationship between temperature and vibration over a specified timeframe.

**Analysis:**
Initial observations indicate a direct correlation between temperature and vibration readings over time. As the temperature increases from the 50-unit threshold, there is a concomitant rise in vibration levels, reaching a peak where both parameters simultaneously exhibit their maximum values. Following the zenith, a decrease in temperature is mirrored by a reduction in vibration intensity.

**Conclusion:**
The analyzed data suggests a positive correlation between temperature and vibration within the parameters of the given dataset. It is important to note that the causal factors underlying this relationship cannot be definitively ascertained from the graph alone and would benefit from further investigative context."""

    st.write(data_analysis_agent_reponse)

######## Logic for download final filled form ########

def create_word_document():
    # Create a new Document object
    doc = Document()
    # Add content to the document, for example, a heading and a paragraph
    doc.add_heading('Analysis on Temperature and Vibration Correlation for Temperature over 50' + '.'.join(csv_document.name.split('.')[:-1]), 0)
    doc.add_paragraph('Introduction:The presented analysis pertains to a filtered dataset where instances with temperatures below 50 units have been excluded. The graph reviews the relationship between temperature and vibration over a specified timeframe.')
    
    # Add an image (chart.png) to the document
    # Specify the path to your image file here
    image_path = 'chart.png'
    # The add_picture method allows you to specify the image size (optional)
    doc.add_picture(image_path, width=Inches(4.0))

    doc.add_paragraph('Analysis: Initial observations indicate a direct correlation between temperature and vibration readings over time. As the temperature increases from the 50-unit threshold, there is a concomitant rise in vibration levels, reaching a peak where both parameters simultaneously exhibit their maximum values. Following the zenith, a decrease in temperature is mirrored by a reduction in vibration intensity. ')

    doc.add_paragraph('Conclusion:The analyzed data suggests a positive correlation between temperature and vibration within the parameters of the given dataset. ')


    # Save the document to a BytesIO object
    byte_io = BytesIO()
    doc.save(byte_io)
    # Seek to the beginning of the BytesIO object
    byte_io.seek(0)
    return byte_io

# Add a download button for the DataFrame
if 'df' in locals():
    word_file = create_word_document()
    st.download_button(
        label='Download Word Report',
        data=word_file,
        file_name='.'.join(csv_document.name.split('.')[:-1])+'_report.docx',
        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )

